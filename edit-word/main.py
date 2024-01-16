import docx

d = docx.Document('docs/demo.docx')

# print(d.paragraphs[3].text)

p = d.paragraphs[3]

# runs divides the text using the bold, italic & underline text fragments as division
print(p.style)

#You can set some variables such as p.runs.underline = True
#You can get the styles of run objects, such as Title, normal, subtitle and certanly you can save them and modify them (p.runs[0].styles = 'Title')

d.add_paragraph("YOOOOOOOOOOOOO")

d.save('docs/demo.docx')